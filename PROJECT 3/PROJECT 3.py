#pip Package Installer for Python
#palindrome using pip
#three bars i.e main menu click settings select projrct file '+' button search 'python-docx' install packages
import tkinter as tk
import os.path
import csv
from tkinter import ttk
from tkinter import messagebox
from docx import Document
import shutil
from tkinter import filedialog

def verf(event) :
    reverse.config(text='')
    if not event.char.isalnum():
        entry.delete(0, tk.END)
        entry.insert(0, entry.get()[::-1])

def pal() :
    given = entry.get()
    if len(given.strip())>0:
        result=given[::-1]
        reverse.config(text=result)
        data(priv.get())
    else:
        messagebox.showwarning('Empty','Given String is Empty')

def clear():
    entry.delete(0,tk.END)
    reverse.config(text='')

def data(user_type):
    if user_type==1:
        rev=entry.get()[::-1]
        if (entry.get()==rev):
            with open(txtf,'a',newline='') as fp:fp.write(entry.get()+'\n')
            doc=Document(docf)
            doc.add_paragraph(entry.get()).bold=True
            doc.save(docf)
        else:
            with open(csvf,'a',newline='') as fp:
                writer=csv.writer(fp)
                writer.writerow([entry.get(),rev])
    elif (user_type==0):
        given = entry.get()
        if len(given.strip()) > 0:
            result = given[::-1]
            reverse.config(text=result)
            if (given.lower() == result.lower()):messagebox.showinfo('Palindrome', 'Is Palindrome')
            else:messagebox.showinfo('Not Palindrome', 'Is Not Palindrome')
        else:messagebox.showwarning('Empty','Inpur is Mandatory')

def downd():
    resdoc=filedialog.asksaveasfile(mode="w",defaultextension=".docx",filetypes=(("DOCX","*.docx"),("All Files","*")))
    if resdoc is None:return
    name=resdoc.name
    basename=os.path.basename(name)
    path=os.path.dirname(name)
    target=os.path.join(path,basename)
    shutil.copyfile(docf,target)

def downc():
    resc=filedialog.asksaveasfile(mode="w",defaultextension=".csv",filetypes=(("CSV","*.csv"),("All Files","*")))
    if resc is None:return
    name=resc.name
    basename=os.path.basename(name)
    path=os.path.dirname(name)
    target=os.path.join(path,basename)
    shutil.copyfile(csvf,target)

def downt():
    rest=filedialog.asksaveasfile(mode="w",defaultextension=".txt",filetypes=(("TXT","*.txt"),("All Files","*")))
    if rest is None:return
    name=rest.name
    basename=os.path.basename(name)
    path=os.path.dirname(name)
    target=os.path.join(path,basename)
    shutil.copyfile(txtf,target)

#<root>main window obj
root= tk.Tk()
root.geometry("1000x400")
root.title("PALINDROME V3")
root.resizable(False,False)
root.config(bg='black')

txtf='PalindromeText.txt'
csvf='ReversedTable.csv'
docf='PalindromeDoc.docx'

if not os.path.exists(txtf):
    with open(txtf,'w',newline='') as fp:
        fp.write('Palindrome List\n============\n')

if not os.path.exists(csvf):
    with open(csvf,'w',newline='') as fp:
        writer=csv.writer(fp)
        writer.writerow(['Given','Reversed'])

if not os.path.exists(docf):
    doc=Document()
    doc.add_heading('Palindrome Check',0)
    doc.save(docf)

my_font=('Times New Roman', 10, 'bold')

priv=tk.IntVar()

head=tk.Label(root,text="Palindrome Check!",font=('Times New Roman', 20, 'bold'),foreground='white',background='black')
head.grid(column=0,row=0,sticky=tk.W,padx=20,pady=20,columnspan=3)

admin=tk.Radiobutton(root,text="Administrator",variable=priv,value=1,font=my_font,activebackground='black',activeforeground='white',command=clear)
admin.grid(column=0,row=1,sticky=tk.W,padx=20,pady=50,columnspan=3)

user=tk.Radiobutton(root,text="User",variable=priv,value=0,font=my_font,activebackground='black',activeforeground='white',command=clear)
user.grid(column=1,row=1,sticky=tk.W,padx=50,pady=50,columnspan=3)

given=ttk.Label(root,text="Given String",font=my_font,foreground='white',background='black')
given.grid(column=0,row=2,sticky=tk.W,padx=20,pady=20)

entry=tk.Entry(root,background='black',foreground='white',font=my_font)
entry.grid(column=1,row=2,sticky=tk.W,padx=20,pady=20,)
entry.bind('<KeyRelease>',verf)

button = tk.Button(root, text='Reverse',width=20,bg='black',font=my_font,fg='white',command=pal)
button.grid(column=2,row=2,sticky=tk.W,padx=20,pady=20)

result=ttk.Label(root,text='Reversed String',foreground='white',background='black',font=my_font)
result.grid(column=0,row=3,sticky=tk.W,padx=20,pady=20)

reverse=ttk.Label(root,text='',background='black',foreground='white',font=my_font)
reverse.grid(column=1,row=3,sticky=tk.W,padx=20,pady=20,columnspan=3)

download=ttk.Label(root,text='Download',foreground='white',background='black',font=my_font)
download.grid(column=0,row=4,sticky=tk.W,padx=20,pady=20)

ddoc=tk.Button(root,text='DOC',width=20,bg='black',font=my_font,fg='white',command=downd)
ddoc.grid(column=1,row=4,sticky=tk.W,padx=20,pady=20)

dcsv=tk.Button(root,text='CSV',width=20,bg='black',font=my_font,fg='white',command=downc)
dcsv.grid(column=2,row=4,sticky=tk.W,padx=20,pady=20)

dct=tk.Button(root,text='TXT',width=20,bg='black',font=my_font,fg='white',command=downt)
dct.grid(column=3,row=4,sticky=tk.W,padx=20,pady=20)

root.mainloop()